import os
import sys
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter as tk
from tkinter import filedialog, messagebox
import random
import logging
import traceback
from PIL import Image, ImageDraw, ImageFont
import tempfile

# Configure logging
def configure_logging():
    log_file = os.path.join(os.getcwd(), 'report_generator.log')
    logging.basicConfig(
        filename=log_file,
        level=logging.DEBUG,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    logging.info("Application started")

class ReportGenerator:
    def __init__(self):
        try:
            logging.info("Initializing application")
            self.root = tk.Tk()
            self.root.title("First Inspection Report Generator")
            self.root.geometry("650x450")
            
            # GUI Elements
            tk.Label(self.root, text="First Inspection Report Generator", 
                     font=("Arial", 16, "bold"), fg="navy").pack(pady=10)
            
            # Input file selection
            file_frame = tk.Frame(self.root)
            file_frame.pack(fill=tk.X, padx=20, pady=5)
            tk.Button(file_frame, text="1. Select Input CSV/Excel File", 
                      command=self.select_input_file, width=25).pack(side=tk.LEFT)
            self.input_file_label = tk.Label(file_frame, text="No file selected", anchor="w")
            self.input_file_label.pack(side=tk.LEFT, padx=10, fill=tk.X, expand=True)
            
            # Images folder selection
            img_frame = tk.Frame(self.root)
            img_frame.pack(fill=tk.X, padx=20, pady=5)
            tk.Button(img_frame, text="2. Select Images Folder", 
                      command=self.select_images_folder, width=25).pack(side=tk.LEFT)
            self.images_folder_label = tk.Label(img_frame, text="No folder selected", anchor="w")
            self.images_folder_label.pack(side=tk.LEFT, padx=10, fill=tk.X, expand=True)
            
            # Output folder selection
            out_frame = tk.Frame(self.root)
            out_frame.pack(fill=tk.X, padx=20, pady=5)
            tk.Button(out_frame, text="3. Select Output Folder", 
                      command=self.select_output_folder, width=25).pack(side=tk.LEFT)
            self.output_folder_label = tk.Label(out_frame, text="No folder selected", anchor="w")
            self.output_folder_label.pack(side=tk.LEFT, padx=10, fill=tk.X, expand=True)
            
            # Generate button
            tk.Button(self.root, text="Generate Reports", command=self.generate_reports,
                      bg="#4CAF50", fg="white", font=("Arial", 12, "bold"), 
                      padx=20, pady=10).pack(pady=20)
            
            # Status label
            self.status_label = tk.Label(self.root, text="Ready to generate reports", 
                                        fg="green", font=("Arial", 10))
            self.status_label.pack(pady=10)
            
            # Instance variables
            self.input_file_path = ""
            self.images_folder_path = ""
            self.output_folder_path = ""
            self.placeholder_cache = {}
            self.header_image_path = None
            self.footer_image_path = None
            
            logging.info("GUI initialized")
            self.root.mainloop()
        except Exception as e:
            logging.exception("Error during initialization")
            messagebox.showerror("Critical Error", f"Initialization failed: {str(e)}\nSee log file for details.")
    
    def select_input_file(self):
        try:
            file_path = filedialog.askopenfilename(
                filetypes=[("Excel/CSV Files", "*.xlsx *.csv"), ("All Files", "*.*")]
            )
            if file_path:
                self.input_file_path = file_path
                self.input_file_label.config(text=os.path.basename(file_path))
                self.status_label.config(text="Input file selected", fg="blue")
        except Exception as e:
            logging.exception("Error selecting input file")
            messagebox.showerror("Error", f"Error selecting input file: {str(e)}")
    
    def select_images_folder(self):
        try:
            folder_path = filedialog.askdirectory()
            if folder_path:
                self.images_folder_path = folder_path
                self.images_folder_label.config(text=os.path.basename(folder_path))
                self.status_label.config(text="Images folder selected", fg="blue")
                
                # Look for header and footer images
                self.find_header_footer_images()
        except Exception as e:
            logging.exception("Error selecting images folder")
            messagebox.showerror("Error", f"Error selecting images folder: {str(e)}")
    
    def find_header_footer_images(self):
        """Find header and footer images in the images folder"""
        self.header_image_path = None
        self.footer_image_path = None
        
        if not self.images_folder_path or not os.path.exists(self.images_folder_path):
            return
            
        for file in os.listdir(self.images_folder_path):
            if file.lower().endswith(('.jpg', '.jpeg', '.png')):
                lower_name = file.lower()
                if 'header' in lower_name:
                    self.header_image_path = os.path.join(self.images_folder_path, file)
                elif 'footer' in lower_name:
                    self.footer_image_path = os.path.join(self.images_folder_path, file)
    
    def select_output_folder(self):
        try:
            folder_path = filedialog.askdirectory()
            if folder_path:
                self.output_folder_path = folder_path
                self.output_folder_label.config(text=os.path.basename(folder_path))
                self.status_label.config(text="Output folder selected", fg="blue")
        except Exception as e:
            logging.exception("Error selecting output folder")
            messagebox.showerror("Error", f"Error selecting output folder: {str(e)}")
    
    def generate_reports(self):
        try:
            if not all([self.input_file_path, self.images_folder_path, self.output_folder_path]):
                messagebox.showerror("Error", "Please select all required files and folders")
                return
            
            self.status_label.config(text="Processing...", fg="orange")
            self.root.update()  # Force UI update
            
            # Read input file
            logging.info(f"Reading input file: {self.input_file_path}")
            if self.input_file_path.endswith('.csv'):
                df = pd.read_csv(self.input_file_path)
            else:
                df = pd.read_excel(self.input_file_path)
            
            # Process each claim
            success_count = 0
            total_count = len(df)
            for idx, row in df.iterrows():
                try:
                    self.status_label.config(text=f"Processing {idx+1}/{total_count}: {row.get('CLAIM #', 'Unknown')}")
                    self.root.update()
                    self.generate_single_report(row)
                    success_count += 1
                except Exception as e:
                    logging.exception(f"Error processing claim: {row.get('CLAIM #', 'Unknown')}")
            
            self.status_label.config(text=f"Generated {success_count}/{total_count} reports", fg="green")
            messagebox.showinfo("Success", f"Successfully generated {success_count} out of {total_count} reports")
        except Exception as e:
            logging.exception("Error generating reports")
            self.status_label.config(text="Error - see log", fg="red")
            messagebox.showerror("Error", f"An error occurred: {str(e)}\nSee log file for details.")
    
    def generate_single_report(self, claim_data):
        # Create document
        doc = Document()
        
        # Add header image if found
        if self.header_image_path:
            try:
                header = doc.sections[0].header
                header.is_linked_to_previous = False
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run()
                run.add_picture(self.header_image_path, width=Inches(6))
                # Add space after header
                doc.add_paragraph("\n\n")
            except Exception as e:
                logging.error(f"Error adding header image: {str(e)}")
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
        
        # Add title
        title = doc.add_heading('FIRST INSPECTION REPORT', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add insured information
        self.add_insured_info(doc, claim_data)
        
        # Add front photo
        self.add_front_photo(doc, claim_data)
        
        # Add Cause of Loss
        self.add_cause_of_loss(doc, claim_data)
        
        # Add Scope of Work
        self.add_scope_of_work(doc, claim_data)
        
        # Add Recommended Reserves
        self.add_recommended_reserves(doc)
        
        # Add Conclusion
        self.add_conclusion(doc)
        
        # Add room photos - now using folder names
        self.add_room_photos_from_folders(doc, claim_data)
        
        # Add footer image if found
        if self.footer_image_path:
            try:
                # Add space before footer
                doc.add_paragraph("\n\n")
                footer = doc.sections[0].footer
                footer.is_linked_to_previous = False
                paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run()
                run.add_picture(self.footer_image_path, width=Inches(6))
            except Exception as e:
                logging.error(f"Error adding footer image: {str(e)}")
        
        # Save document
        filename = f"FIRST INSPECTION REPORT - CLAIM# {claim_data.get('CLAIM #', 'PR0000')} - {claim_data.get('INSURED/POLICYHOLDER', 'UNKNOWN').split()[0].upper()} - {claim_data.get('ADDRESS', 'UNKNOWN').replace(',', '').replace(' ', '_')}.docx"
        save_path = os.path.join(self.output_folder_path, filename)
        doc.save(save_path)
        logging.info(f"Saved report: {save_path}")
    
    def add_insured_info(self, doc, claim_data):
        doc.add_paragraph(f"INSURED/POLICYHOLDER: {claim_data.get('INSURED/POLICYHOLDER', 'Unknown')}")
        doc.add_paragraph(f"ADDRESS: {claim_data.get('ADDRESS', 'Unknown')}")
        doc.add_paragraph(f"INSURER: {claim_data.get('INSURER', 'Unknown')}")
        doc.add_paragraph(f"CLAIM #: {claim_data.get('CLAIM #', 'PR0000')}")
        doc.add_paragraph(f"ADJUSTER/ CLAIM REP: {claim_data.get('ADJUSTER/ CLAIM REP', 'Unknown')}")
        
        # Format dates consistently
        inspection_date = claim_data.get('DATE OF INSPECTION', 'Unknown')
        loss_date = claim_data.get('DATE OF LOSS', 'Unknown')
        report_date = claim_data.get('DATE OF REPORT', 'Unknown')
        
        doc.add_paragraph(f"DATE OF INSPECTION: {inspection_date}")
        doc.add_paragraph(f"DATE OF LOSS: {loss_date}")
        doc.add_paragraph(f"DATE OF REPORT: {report_date}")
        doc.add_paragraph(f"TYPE OF LOSS: {claim_data.get('TYPE OF LOSS', 'Unknown')}")
        
        doc.add_paragraph()  # Add empty line
    
    def add_front_photo(self, doc, claim_data):
        # Try to find front photo
        front_photo_path = self.find_photo('front', 'exterior', 'house')
        if front_photo_path:
            doc.add_paragraph("Front Photo:")
            doc.add_picture(front_photo_path, width=Inches(3.25), height=Inches(2.25))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph("Image 1")
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            # Create placeholder if no front photo found
            placeholder = self.create_placeholder_image("Front of House", claim_data.get('ADDRESS', 'Unknown'))
            if placeholder:
                doc.add_paragraph("Front Photo:")
                doc.add_picture(placeholder, width=Inches(3.25), height=Inches(2.25))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph("Image 1 - Placeholder")
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                doc.add_paragraph("Front Photo: [Image not available]")
        
        doc.add_paragraph()  # Add empty line
    
    def add_cause_of_loss(self, doc, claim_data):
        doc.add_paragraph("CAUSE OF LOSS:", style='Heading 2')
        cause = claim_data.get('CAUSE OF LOSS', 'Unknown cause of loss')
        doc.add_paragraph(cause)
        doc.add_paragraph()  # Add empty line
    
    def add_scope_of_work(self, doc, claim_data):
        doc.add_paragraph("SCOPE OF WORK:", style='Heading 2')
        doc.add_paragraph("The following is a brief outline of the work to be completed on the contents portion of this claim.")
        doc.add_paragraph()
        
        scope = claim_data.get('SCOPE OF WORK', '')
        if isinstance(scope, str):
            # Split by <br> tags or numbers
            items = [item.strip() for item in scope.split('<br>') if item.strip()]
            if len(items) <= 1:  # If not split by <br>, try numbered items
                items = [item.strip() for item in scope.split('\n') if item.strip()]
            
            for item in items:
                # Clean up item (remove numbers if they exist)
                clean_item = item.split('.', 1)[-1].strip() if '.' in item else item
                doc.add_paragraph(f"• {clean_item}", style='List Bullet')
        else:
            doc.add_paragraph("• Scope of work details not available", style='List Bullet')
        
        doc.add_paragraph()  # Add empty line
    
    def add_recommended_reserves(self, doc):
        doc.add_paragraph("RECOMMENDED RESERVES FOR TRINITY'S INVOLVEMENT:", style='Heading 2')
        
        # Generate random but realistic amounts
        indemnity = random.randint(15000, 30000)
        pricing_expense = random.randint(3000, 6000)
        total_replacement = random.randint(3000, 10000)
        
        doc.add_paragraph(f"The estimated cost for Trinity's involvement is as follows:")
        doc.add_paragraph(f"• Indemnity Work: Should not exceed ${indemnity:,.2f} plus HST")
        doc.add_paragraph("Our actual cost will be adjusted once the exact scope of approved work is known. The recommended estimate is only based on visual inspection for reserves setting purposes.")
        doc.add_paragraph()
        doc.add_paragraph(f"• Trinity Listing & Pricing Expense Reserve: Should not exceed ${pricing_expense:,.2f} plus HST")
        doc.add_paragraph()
        
        doc.add_paragraph("RECOMMENDED RESERVES FOR THE TOTAL CONTENTS LOSS:", style='Heading 2')
        doc.add_paragraph(f"Based on a visual inspection of the extent of non-salvageable items on the main floor, we believe that the total replacement cost for the non-salvageable items should not exceed ${total_replacement:,.2f} plus HST.")
        doc.add_paragraph()  # Add empty line
    
    def add_conclusion(self, doc):
        doc.add_paragraph("CONCLUSION:", style='Heading 2')
        doc.add_paragraph("Once our scope of work is approved, we can attend and begin the pack out process.")
        doc.add_paragraph()
        doc.add_paragraph("Thank You,")
        doc.add_paragraph()
        doc.add_paragraph("Mo Waez")
        doc.add_paragraph("Trinity Contents Management")
        doc.add_paragraph("mo@trinitycontents.com")
        doc.add_paragraph("(647) 613-2246")
        doc.add_paragraph()  # Add empty line
    
    def add_room_photos_from_folders(self, doc, claim_data):
        """Add photos organized by room folders"""
        if not self.images_folder_path or not os.path.exists(self.images_folder_path):
            return
            
        # Get all room folders
        room_folders = []
        for item in os.listdir(self.images_folder_path):
            item_path = os.path.join(self.images_folder_path, item)
            if os.path.isdir(item_path):
                room_folders.append(item)
        
        # If no room folders found, use default rooms
        if not room_folders:
            room_folders = [
                'kitchen', 'dining', 'living', 'bedroom1', 'bedroom2', 
                'bathroom', 'storage', 'basement', 'garage'
            ]
        
        # Image counter starts at 2 (front photo is 1)
        image_counter = 2
        
        for room in room_folders:
            room_path = os.path.join(self.images_folder_path, room)
            room_photos = []
            
            # Get all valid image files from room folder
            if os.path.exists(room_path) and os.path.isdir(room_path):
                for file in os.listdir(room_path):
                    if file.lower().endswith(('.jpg', '.jpeg', '.png')):
                        room_photos.append(os.path.join(room_path, file))
            
            # If no photos found, create a placeholder
            if not room_photos:
                placeholder = self.create_placeholder_image(
                    f"{room.capitalize()} Area", 
                    claim_data.get('ADDRESS', 'Unknown')
                )
                if placeholder:
                    room_photos = [placeholder]
            
            if room_photos:
                # Add room heading
                room_name = ' '.join(word.capitalize() for word in room.replace('_', ' ').split())
                doc.add_paragraph(f"{room_name.upper()} AREA", style='Heading 2')
                
                # Create table for photos (2 columns)
                table = doc.add_table(rows=0, cols=2)
                table.autofit = False
                
                # Set column widths
                for col in table.columns:
                    col.width = Inches(3.5)
                
                # Add photos to table
                row = None
                for i, photo_path in enumerate(room_photos[:4]):  # Max 4 photos per room
                    if i % 2 == 0:
                        row = table.add_row()
                    
                    cell = row.cells[i % 2]
                    try:
                        # Add image to cell
                        cell_paragraph = cell.paragraphs[0]
                        run = cell_paragraph.add_run()
                        run.add_picture(photo_path, width=Inches(3.25), height=Inches(2.25))
                        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Add image number
                        img_text = f"Image {image_counter}"
                        if "placeholder" in photo_path.lower():
                            img_text += " - Placeholder"
                        cell_paragraph = cell.add_paragraph(img_text)
                        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        image_counter += 1
                    except Exception as e:
                        logging.error(f"Error adding image {photo_path}: {str(e)}")
                        cell.text = f"Image not available\n{os.path.basename(photo_path)}"
                
                doc.add_paragraph()  # Add empty line
    
    def find_photo(self, *keywords):
        """Find a photo containing any of the keywords in its name"""
        if not self.images_folder_path or not os.path.exists(self.images_folder_path):
            return None
            
        for file in os.listdir(self.images_folder_path):
            if file.lower().endswith(('.jpg', '.jpeg', '.png')):
                if any(keyword.lower() in file.lower() for keyword in keywords):
                    return os.path.join(self.images_folder_path, file)
        return None
    
    def create_placeholder_image(self, title, subtitle):
        """Create a placeholder image with text"""
        try:
            # Check cache
            cache_key = f"{title}_{subtitle}"
            if cache_key in self.placeholder_cache:
                return self.placeholder_cache[cache_key]
                
            # Create image
            width, height = 800, 600
            img = Image.new('RGB', (width, height), color=(230, 230, 230))
            draw = ImageDraw.Draw(img)
            
            # Add border
            draw.rectangle([(10, 10), (width-10, height-10)], outline=(180, 180, 180), width=3)
            
            # Add title
            try:
                title_font = ImageFont.truetype("arialbd.ttf", 40)
            except:
                title_font = ImageFont.load_default()
            
            title_width = draw.textlength(title, font=title_font) if hasattr(draw, 'textlength') else 400
            draw.text(((width - title_width) // 2, height // 3), title, 
                      fill=(100, 100, 100), font=title_font)
            
            # Add subtitle
            try:
                subtitle_font = ImageFont.truetype("arial.ttf", 30)
            except:
                subtitle_font = ImageFont.load_default()
            
            subtitle_width = draw.textlength(subtitle, font=subtitle_font) if hasattr(draw, 'textlength') else 400
            draw.text(((width - subtitle_width) // 2, height // 2), subtitle, 
                      fill=(150, 150, 150), font=subtitle_font)
            
            # Add camera icon
            draw.ellipse([(width//2-40, height//1.7-40), (width//2+40, height//1.7+40)], 
                         outline=(180, 180, 180), width=3)
            draw.line([(width//2-25, height//1.7-25), (width//2+25, height//1.7+25)], 
                      fill=(180, 180, 180), width=3)
            draw.line([(width//2-25, height//1.7+25), (width//2+25, height//1.7-25)], 
                      fill=(180, 180, 180), width=3)
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(suffix='.jpg', delete=False)
            img.save(temp_file.name)
            temp_file.close()
            
            # Add to cache
            self.placeholder_cache[cache_key] = temp_file.name
            return temp_file.name
        except Exception as e:
            logging.error(f"Error creating placeholder: {str(e)}")
            return None

if __name__ == "__main__":
    try:
        configure_logging()
        ReportGenerator()
    except Exception as e:
        error_msg = f"Unhandled error: {str(e)}\n{traceback.format_exc()}"
        logging.exception(error_msg)
        
        # Try to show error message even if Tk fails
        try:
            root = tk.Tk()
            root.withdraw()
            messagebox.showerror("Critical Error", "Application failed to start.\nSee log file for details.")
        except:
            print(error_msg)
            input("Press Enter to exit...")